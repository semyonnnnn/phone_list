from fastapi import APIRouter, Request, HTTPException
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io

router = APIRouter()

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_table_borders(table, size=4, color="000000"):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), str(size))   # border thickness in eighths of a point
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), color)
        borders.append(edge_el)
    tblPr.append(borders)

# Row-level shading for hierarchy distinction
BOSS_BG = "D9D9D9"       # darker gray for bosses
MINIBOSS_BG = "F2F2F2"   # lighter gray for miniboses
REGULAR_BG = None        # no fill for regular employees

def sort_key(emp):
    """Bosses first (0), miniboses second (1), everyone else last (2).
    Within each tier, keep original relative order (stable sort)."""
    if emp.get("isBoss"):
        return 0
    if emp.get("isMiniBoss"):
        return 1
    return 2

@router.post("/api/download")
async def generate_word_document(request: Request):
    try:
        body = await request.json()
        phones = body.get("phones", [])

        # Group records by department group manually
        departments = {}
        for p in phones:
            group_name = p.get("group") or "Отдел"
            if group_name not in departments:
                departments[group_name] = []
            departments[group_name].append(p)

        doc = Document()

        # Set landscape orientation using the correct WD_ORIENT enum
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.0)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

        col_widths = [Inches(4.5), Inches(1.5), Inches(1.2), Inches(1.0), Inches(0.8)]

        for group_name, emp_list in departments.items():
            table = doc.add_table(rows=0, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(table)

            # 1. Department Header Row (Spanning all 5 columns)
            header_row = table.add_row()
            cell_dep = header_row.cells[0]
            for i in range(1, 5):
                cell_dep.merge(header_row.cells[i])

            set_cell_background(cell_dep, "224376")
            p = cell_dep.paragraphs[0]
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(group_name)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(255, 255, 255)

            # 2. Table Column Headers
            titles_row = table.add_row()
            titles = ["ФИО", "Телефон", "Добавочный", "Кабинет", "IP"]
            for idx, title_text in enumerate(titles):
                cell = titles_row.cells[idx]
                cell.width = col_widths[idx]
                set_cell_background(cell, "2B579A")

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(title_text)
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)

            # 3. Employee Data Rows — bosses first, then miniboses, then the rest
            sorted_emp_list = sorted(emp_list, key=sort_key)

            for emp in sorted_emp_list:
                row = table.add_row()
                row_data = [
                    emp.get("person", ""),
                    emp.get("phone", ""),
                    emp.get("extension", ""),
                    emp.get("cabinet", ""),
                    emp.get("ip", "")
                ]

                if emp.get("isBoss"):
                    row_bg = BOSS_BG
                elif emp.get("isMiniBoss"):
                    row_bg = MINIBOSS_BG
                else:
                    row_bg = REGULAR_BG

                for idx, val in enumerate(row_data):
                    cell = row.cells[idx]
                    cell.width = col_widths[idx]
                    if row_bg:
                        set_cell_background(cell, row_bg)

                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    run = p.add_run(str(val))
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    # Slightly emphasize boss/miniboss names for extra clarity
                    if emp.get("isBoss"):
                        run.bold = True

        # Save document to an in-memory buffer
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Define an explicit chunk generator to stream the binary safely
        def iterfile():
            chunk_size = 1024 * 64
            while chunk := file_stream.read(chunk_size):
                yield chunk

        return StreamingResponse(
            iterfile(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=phone_directory.docx"}
        )

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Word generation error: {str(e)}"
        )